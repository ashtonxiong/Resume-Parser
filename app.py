import requests
from bs4 import BeautifulSoup
import re
from docx import Document
import os
from urllib.parse import urlparse


job_title = ""
company = ""


def extract_intern_jobs(url):
    response = requests.get(url)
    data = response.text
    soup = BeautifulSoup(data, "html.parser")

    # Customize these patterns based on the HTML structure of the websites you're working with
    title_patterns = [
        # Case-insensitive search for "intern"
        re.compile(r'\bintern\b', flags=re.IGNORECASE),
        # Add more patterns if needed, tailored to the specific HTML structure of different websites
    ]

    job_names = []

    # Extract job names using different patterns
    for pattern in title_patterns:
        matching_elements = soup.find_all(
            name=lambda tag: tag and pattern.search(str(tag)))
        for element in matching_elements:
            job_names.append(element.get_text(strip=True))

    return job_names


def jobname(intern_jobs):
    obj = intern_jobs[len(intern_jobs) - 1]
    arr = obj.split(',')

    for text in arr:
        temp = text.split(':')
        match = re.search(r'"title":"([^"]*)"', text)
        match2 = temp[0].strip()

        if "title" in match2:
            return temp[1]

        if match and "title" in temp[0]:
            return temp[1]


def company_name(intern_jobs):
    obj = intern_jobs[len(intern_jobs) - 1]
    arr = obj.split(',')

    for text in arr:
        temp = text.split(':')
        # print(temp)
        if "mapQueryLocation" in temp:
            print(temp)
        # match = re.search(r'"mapQueryLocation"([^"]*)"', text)

        # if match and "title" in temp[0]:
        #     return temp[1]


def extract_domain_from_url(url):
    parsed_url = urlparse(url)
    domain = parsed_url.netloc
    arr = domain.split('.')
    if "myworkdayjobs" in domain:
        domain = arr[0]
    else:
        domain = arr[1]
    return domain


def load_word_doc(doc_path):
    doc = Document(doc_path)
    return doc


def add_inputs_to_doc(doc, inputs):
    for input_key, input_value in inputs.items():
        # Assuming placeholders in the Word document are enclosed in curly braces, like {input_key}
        for paragraph in doc.paragraphs:
            # arr = paragraph.text.split('\n')
            # print(arr)
            if "{" + input_key + "}" in paragraph.text:
                paragraph.text = paragraph.text.replace(
                    "{" + input_key + "}", str(input_value))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "{" + input_key + "}" in cell.text:
                        cell.text = cell.text.replace(
                            "{" + input_key + "}", str(input_value))


def save_word_doc(doc, output_path):
    doc.save(output_path)


def main():
    # Example usage for the provided URL

    URL = input()

    intern_jobs = extract_intern_jobs(URL)
    job_title = jobname(intern_jobs)
    job_title = job_title[2:len(job_title) - 1]
    company = extract_domain_from_url(URL)
    company = company[0].upper() + company[1:]
    combine = company + " " + job_title

    word_doc_path = "/Users/ashtonxiong/Desktop/Resume:CV's/Cover Letters/A. Xiong Cover Letter.docx"
    loaded_doc = load_word_doc(word_doc_path)

    inputs = {"NAME OF COMPANY": company, "ROLE NAME": job_title,
              "NAME OF COMPANY AND ROLE": combine}

    add_inputs_to_doc(loaded_doc, inputs)

    new_doc_name = "A. Xiong Cover Letter " + company
    # Replace spaces with underscores
    new_doc_name = new_doc_name.replace(" ", "_")

    output_path = "/Users/ashtonxiong/Desktop/Resume:CV's/Cover Letters/" + \
        new_doc_name + ".docx"
    save_word_doc(loaded_doc, output_path)


main()
